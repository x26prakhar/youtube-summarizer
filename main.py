import os
from datetime import datetime
from pathlib import Path

from fastapi import FastAPI, Request, HTTPException
from fastapi.responses import HTMLResponse, FileResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from pydantic import BaseModel
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from transcript import get_clean_transcript
from summarizer import process_transcript

app = FastAPI(title="YouTube Transcript Summarizer")

# Setup templates and static files
BASE_DIR = Path(__file__).resolve().parent
app.mount("/static", StaticFiles(directory=BASE_DIR / "static"), name="static")
templates = Jinja2Templates(directory=BASE_DIR / "templates")

# Ensure outputs directory exists
OUTPUTS_DIR = BASE_DIR / "outputs"
OUTPUTS_DIR.mkdir(exist_ok=True)


class ProcessRequest(BaseModel):
    url: str


class ProcessResponse(BaseModel):
    video_id: str
    summary: str
    notes: str
    filename: str
    docx_filename: str


def create_word_document(video_id: str, url: str, takeaways: str, clean_transcript: str, filepath: Path):
    """Create a Word document with the results."""
    doc = Document()

    # Title
    title = doc.add_heading('YouTube Video Notes', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Video info
    doc.add_paragraph(f"Video ID: {video_id}")
    doc.add_paragraph(f"URL: https://youtube.com/watch?v={video_id}")
    doc.add_paragraph(f"Processed: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    doc.add_paragraph()  # Spacer

    # Key Takeaways
    doc.add_heading('Key Takeaways', level=1)
    doc.add_paragraph(takeaways)

    doc.add_paragraph()  # Spacer

    # Clean Transcript
    doc.add_heading('Clean Transcript', level=1)
    doc.add_paragraph(clean_transcript)

    doc.save(filepath)


@app.get("/", response_class=HTMLResponse)
async def home(request: Request):
    """Serve the main web interface."""
    return templates.TemplateResponse("index.html", {"request": request})


@app.post("/process", response_model=ProcessResponse)
async def process_video(data: ProcessRequest):
    """Process a YouTube video URL and return summary + notes."""
    try:
        # Extract and clean transcript
        print(f"\n{'='*50}")
        print(f"Processing: {data.url}")
        print(f"{'='*50}")

        video_id, transcript = get_clean_transcript(data.url)
        print(f"Video ID: {video_id}")
        print(f"Transcript length: {len(transcript)} characters")

        # Format transcript (no AI - fast mode)
        print("Formatting transcript...")
        result = process_transcript(transcript)
        print("Done!")

        # Save files
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        md_filename = f"{video_id}_{timestamp}.md"
        docx_filename = f"{video_id}_{timestamp}.docx"
        md_filepath = OUTPUTS_DIR / md_filename
        docx_filepath = OUTPUTS_DIR / docx_filename

        # Save markdown
        markdown_content = f"""# YouTube Video Notes

**Video ID:** {video_id}
**URL:** https://youtube.com/watch?v={video_id}
**Processed:** {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}

---

## Key Takeaways

{result['summary']}

---

## Clean Transcript

{result['notes']}
"""
        md_filepath.write_text(markdown_content, encoding="utf-8")

        # Save Word document
        create_word_document(video_id, data.url, result['summary'], result['notes'], docx_filepath)

        print(f"Saved to: {md_filepath}")
        print(f"Saved to: {docx_filepath}")

        # Print to console
        print(f"\n{'='*50}")
        print("KEY TAKEAWAYS:")
        print(f"{'='*50}")
        print(result['summary'])
        print(f"{'='*50}\n")

        return ProcessResponse(
            video_id=video_id,
            summary=result['summary'],
            notes=result['notes'],
            filename=md_filename,
            docx_filename=docx_filename
        )

    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        print(f"Error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/download/{filename}")
async def download_file(filename: str):
    """Download the generated file."""
    filepath = OUTPUTS_DIR / filename
    if not filepath.exists():
        raise HTTPException(status_code=404, detail="File not found")

    media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document" if filename.endswith('.docx') else "text/markdown"

    return FileResponse(
        path=filepath,
        filename=filename,
        media_type=media_type
    )


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
